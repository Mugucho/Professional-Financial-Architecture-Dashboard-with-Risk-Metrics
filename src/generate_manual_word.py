from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def crear_manual_word():
    doc = Document()

    # --- ESTILO DE TÍTULOS ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # --- PORTADA ESTILO TESIS ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDAD DEL TRADING ALGORÍTMICO\n")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("\n" * 5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "MARKET ARCHITECT PRO:\nSISTEMA INTEGRADO DE ANÁLISIS TÉCNICO, MACROECONÓMICO Y GESTIÓN DE RIESGO\n"
    )
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph("\n" * 5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Autor: Miguel Calle Romero\nLocalización: Trenton, New Jersey\nFecha: Abril 2026"
    )
    run.font.size = Pt(12)

    doc.add_page_break()

    # --- CONTENIDO ---
    doc.add_heading("CAPÍTULO I: INTRODUCCIÓN Y ARQUITECTURA", 1)
    doc.add_paragraph(
        "Market Architect Pro representa una evolución en los sistemas de soporte de decisiones financieras. "
        "Inspirado en la arquitectura de control de vehículos no tripulados (drones), el sistema opera mediante "
        "una estructura modular donde cada componente procesa una capa específica de información (técnica, "
        "psicológica, algorítmica y macroeconómica)."
    )

    doc.add_heading("CAPÍTULO II: MARCO TEÓRICO (CONCEPTOS)", 1)
    conceptos = [
        (
            "RSI (Relative Strength Index):",
            "Indicador de oscilación que evalúa la fuerza relativa del precio.",
        ),
        (
            "Patrón IHS (Inverted Head and Shoulders):",
            "Formación técnica que señala el agotamiento de una tendencia bajista y el inicio de un ciclo alcista.",
        ),
        (
            "Análisis de Sentimiento (NLP):",
            "Técnica de Inteligencia Artificial que traduce lenguaje humano (noticias) en datos cuantitativos de polaridad.",
        ),
        (
            "Backtesting:",
            "Validación estadística de una tesis de inversión sobre datos históricos.",
        ),
    ]
    for titulo, desc in conceptos:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(titulo)
        run.bold = True
        p.add_run(f" {desc}")

    doc.add_heading("CAPÍTULO III: ESPECIFICACIONES TÉCNICAS (BACKEND)", 1)
    modulos = [
        (
            "src/risk_management.py",
            "Gestión de seguridad mediante el 'Risk Pipeline' (Drawdown, Exposure, Reconciliation).",
        ),
        (
            "src/pattern_recognition.py",
            "Detección algorítmica de 60+ patrones mediante TA-Lib y lógica personalizada.",
        ),
        (
            "src/oracle.py",
            "Motor de procesamiento de noticias en tiempo real con integración de NewsAPI.",
        ),
        (
            "src/macro_analysis.py",
            "Sincronizador de hitos históricos y eventos geopolíticos.",
        ),
    ]
    for mod, desc in modulos:
        doc.add_heading(mod, level=2)
        doc.add_paragraph(desc)

    doc.add_heading("CAPÍTULO IV: MANUAL DE OPERACIÓN", 1)
    doc.add_paragraph(
        "El acceso al sistema se realiza a través de una interfaz de Streamlit dividida en siete (7) pestañas operativas:\n"
    )
    pestañas = [
        "1. Análisis Técnico: Gráficos de velas y niveles críticos.",
        "2. Patrones: Identificación de señales psicológicas.",
        "3. Riesgo y Volumen: Diagnóstico de salud del mercado.",
        "4. Trading Bot: Ejecución automatizada con Alpaca.",
        "5. Datos: Auditoría de Dataframes procesados.",
        "6. Backtesting: Simulación de rentabilidad.",
        "7. Macro Impact: Correlación entre noticias y precio.",
    ]
    for item in pestañas:
        doc.add_paragraph(item, style="List Number")

    doc.add_page_break()
    doc.add_paragraph("Fin del Documento Técnico.").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    # --- GUARDAR ---
    doc.save("Market_Architect_Pro_Manual.docx")
    print("✅ Documento Word generado con éxito: Market_Architect_Pro_Manual.docx")


if __name__ == "__main__":
    crear_manual_word()
